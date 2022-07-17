# -*- coding:utf-8 -*-
# Install mtranslate, googletrans for translations
# Install python-docx for managing the Word Files.
# Install Pandas to manage the Excel file and bring the information
# Import Shutil to remove the directory

import os
from deep_translator import GoogleTranslator  
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
import pandas as pd
from shutil import rmtree


def importSouthData():
    # Define the path for the excel file
    excelPath = os.getcwd()
    excelPath = os.path.join(excelPath + "\GaN\Activity_Guide_Changes\GaN_cons_and_dates.xlsx")

    # Get Data from the Excel File using Pandas
    # Capitalize  constellations names for a later comparison
    dfSouth = pd.read_excel (excelPath,sheet_name='South')
    dfSouth['Constellations'] = dfSouth['Constellations'].str.capitalize()

    #store constellation and date information in respective variables
    southCons = dfSouth['Constellations']
    southDates = dfSouth['Dates']

    #updates the constellation: creates a variable to hold the new constellation, uses the old constellation
    #information to find the new constellation
    newConstellationSouth = {}
    
    # Ading key and values to the new dictionary
    for i in range(len(southCons)):
        newConstellationSouth[southCons[i]] = southDates[i] 

    #return the dictionary with the North Data
    return newConstellationSouth

def createSouthDir(year, constellations):
    cons = constellations
    year = year
    savePath = os.getcwd() 
    savePath = os.path.join(savePath + "\GaN\docs_changed")
    #rmtree(savePath)
    #os.mkdir(savePath)
    paths = []
    for con in cons:
        savePath = os.getcwd() 
        savePath = os.path.join(savePath + "\GaN\docs_changed\GaN_South_{year}_ActivityGuide_{con}".format(year = year, con = con))        
        os.mkdir(savePath)
        paths.append(savePath)
    
    return paths

def createSouthPaths(directories, languages, latitudes):

    dirPaths = []
    for lang in languages:
        for direc in directories:
            for lat in latitudes:
                dirPaths.append(direc + "_" + lat + "_" + lang)
    return dirPaths


#opens document that will be edited
def openWordDoc2(filename):
    document = Document(filename)
    return document


#updating southern hemisphere information (constellation, date, text displayed to user)
#######################################################################################
########################All the information that needs to change#######################
#######################################################################################

def southTranslation(dirPaths):
    #updating southern hemisphere information (constellation, date, text displayed to user)
    southConstellationReplacement = {
        "ChileanSpanish" : "Escorpio",
        "English" : "Scorpius",
        "French" : "Scorpion",
        "Indonesian" : "Scorpio",
        "Portuguese" : "Escorpião",
        "Spanish" : "Escorpio"
        }
        
    southDateReplacement = {
        "ChileanSpanish" : "Del 4 al 13 de julio y del 2 al 11 de agosto",
        "English" : "July 4-13 and August 2-11",
        "French" : "Du 4 au 13 juillet et du 2 au 11 août",
        "Indonesian" : "4 Juli - 13 Juli dan 2 Agustus - 11 Agustus",
        "Portuguese" : "4 a 13 de julho e 2 a 11 de agosto.",
        "Spanish" : "del 4 al 13 de julio y del 2 al 11 de agosto"
        }

    southHeadingFirst = {
        "ChileanSpanish" : "",
        "English" : " Campaign Dates that use the ",
        "French" : "Dates de la campagne ",
        "Indonesian" : "Waktu Kampanye ",
        "Portuguese" : "Datas das campanhas de ",
        "Spanish" : "Fechas de la campaña año "
        }

    southHeadingMiddle = {
        "ChileanSpanish" : " Fechas de campaña para la constelación del ",
        "English" : "",
        "French" : " qui utilisent la ",
        "Indonesian" : " yang menggunakan ",
        "Portuguese" : " que usam a ",
        "Spanish" : " que utilizan la "
        }

    southHeadingLast = {
        "ChileanSpanish" : ": ",
        "English" : ": ",
        "French" : ": ",
        "Indonesian" : ": ",
        "Portuguese" : ": ",
        "Spanish" : ": "
        }

    firstParagraphfirst = {
        "ChileanSpanish" : "Usted está participando en una campaña mundial para observar y registrar las estrellas visibles más débiles como un medio para medir la contaminación lumínica en un lugar determinado. Localizando y observando la ",
        "English" : "You are participating in a global campaign to observe and record the faintest stars visible as a means of measuring light pollution in a given location. By locating and observing the constellation ",
        "French" :   "Vous allez participer à une campagne mondiale d’observation pour détecter les plus faibles étoiles visibles afin de mesurer la pollution lumineuse sur un site donné. Partout dans le monde, en localisant et en observant la ",
        "Indonesian" : "Anda sedang berpartisipasi dalam kampanye global pengamatan dan pencatatan penampakan bintang paling redup untuk pengukuran tingkat polusi cahaya di suatu lokasi. Melalui pengamatan dan identifikasi  ",
        "Portuguese" : "Está a participar numa campanha global para observar e registar as estrelas mais fracas visíveis como forma de medir a poluição luminosa num determinado local. Localizando e observando a  ",
        "Spanish" : "Usted está participando en una campaña mundial para observar y registrar las estrellas visibles más débiles como un medio para medir la contaminación lumínica en un lugar determinado. Localizando y observando la  ",

    }

    firstParagraphLast = {
        "ChileanSpanish" : " en el cielo nocturno y comparándolo con las cartas estelares, la gente de todo el mundo aprenderá cómo las luces de su comunidad contribuyen a la contaminación lumínica. Sus contribuciones a la base de datos en línea documentarán el cielo nocturno visible.",
        "English" : " in the night sky and comparing it to stellar charts, people from around the world will learn how the lights in their community contribute to light pollution. Your contributions to the online database will document the visible nighttime sky.",
        "French" :   " dans le ciel nocturne et en la comparant aux cartes stellaires, les participants, apprendront comment l’éclairage, dans leur environnement local, influence la pollution lumineuse. Vos contributions à la base de données en ligne permettront de mesurer la qualité du ciel nocturne.",
        "Indonesian" : " di langit malam dan membandingkannya dengan peta bintang, masyarakat di seluruh dunia dapat mengetahui dan mempelajari seberapa besar kontribusi cahaya di lingkungannya terhadap polusi cahaya. Kontribusi data anda pada basis data online akan membantu mendokumentasikan langit malam yang tampak di berbagai lokasi.",
        "Portuguese" : " no céu noturno e,  comparando-a com cartas estelares, pessoas de todo o mundo aprenderão  como as luzes da sua comunidade contribuem para a poluição luminosa. As suas contribuições para a base de dados on-line irão documentar a visibilidade do céu noturno em todo o mundo.",
        "Spanish" : " en el cielo nocturno y comparándolo con las cartas estelares, la gente de todo el mundo aprenderán cómo las luces de su comunidad contribuyen a la contaminación lumínica. Sus contribuciones a la base de datos en línea documentarán el cielo nocturno visible.",

    }

    ##################################################################################################
    ##################################################################################################
            ###	End of the changes section defining things that need to be changed			###
    ##################################################################################################
    ##################################################################################################

    # Get data from the Excel file and bring the created Paths
    dirPath = dirPaths
    southData = importSouthData()


# Organize the Languages by lists to make better translations
    CountryList1 = ("Chilean_Spanish", "French", "Indonesian", "Portuguese","Spanish")
    CountryList2 = ("English")
    

    # Getting data from the Paths
    languageBase = dirPath.split('_')[-1]
    latitude = dirPath.split('_')[-2]
    constName = dirPath.split('_')[-3]
    year = dirPath.split('_')[-5]
    #thaiYear = int(year)+ 543

    #Be sure to change the websites into the word files
    website1 = "astro/maps/GaNight/2018/"
    website2 = "astro/maps/GaNight/2019/"

    
    # Define the Word file path as the original file
    wordPath = os.path.abspath("..\Gan\GaN\docs_to_change\GaN2018_ActivityGuide_Scorpius_S_")
    workingDoc = openWordDoc2(wordPath + str(languageBase) + ".docx")

    # styles of each paragraph to kkep the original word styles
    objStyles = workingDoc.styles
    objCharstyle = objStyles.add_style('GaNStyle', WD_STYLE_TYPE.CHARACTER)
    objFont = objCharstyle.font
    objFont.name = 'Calibri'
    objFont.size = Pt(14)
    
    objStyles2 = workingDoc.styles
    objCharstyle2 = objStyles2.add_style('GaNParagraph', WD_STYLE_TYPE.CHARACTER)
    objFont2 = objCharstyle2.font
    objFont2.name = 'Calibri'
    objFont2.size = Pt(10)
    
    objStyles3 = workingDoc.styles
    objCharstyle3 = objStyles3.add_style('GaNLinks', WD_STYLE_TYPE.CHARACTER)
    objFont3 = objCharstyle3.font
    objFont3.name = 'Calibri'
    objFont3.size = Pt(9.5)
    objFont3.bold = True
    objFont3.underline = True
    objFont3.color.rgb = RGBColor(0,0,128)


    #Define the base language in deep_translator and translate it into de destiny language
    if constName != "Canis Major" :
        constellationTranslated =GoogleTranslator(source ='english', target = languageBase.lower()).translate(constName +" constellation")
        dateTranslated = GoogleTranslator(source ='english', target = languageBase.lower()).translate(southData.get(constName))
    else:
        constellationTranslated = "Canis Major"
        dateTranslated = southData.get('Canis major')


    # Replace the translations in the proper places
    for languageSelected, date in southDateReplacement.items():
        if languageSelected == languageBase:
            for paragraph in workingDoc.paragraphs:
                #If the contellation's name is in the paragraph, delete the paragraph and add a new one with the translations
                if southConstellationReplacement[languageBase] in paragraph.text:
                    # Replace only if the name and the date is on the paragraph, organizng with the grammar of each language
                    if date in paragraph.text:
                        paragraph.clear()
                        if languageBase in CountryList1:
                            paragraph.add_run(southHeadingFirst[languageBase]+ str(year) + southHeadingMiddle[languageBase]+ constellationTranslated + southHeadingLast[languageBase] + dateTranslated + ".", style = 'GaNStyle')
                        elif languageBase in CountryList2:
                            paragraph.add_run(str(year) + southHeadingFirst[languageBase]+ constellationTranslated + southHeadingMiddle[languageBase] + southHeadingLast[languageBase]+ dateTranslated + ".", style = 'GaNStyle' )
                    
                    # Replace only if the constellation's name is in the paragraph
                    else:
                        paragraph.clear()
                        if(languageBase!= 'Japanese'):
                            paragraph.add_run(firstParagraphfirst[languageBase] + constellationTranslated + firstParagraphLast[languageBase], style = 'GaNParagraph')
                        else:
                            paragraph.add_run(firstParagraphfirst[languageBase] + firstParagraphLast[languageBase] + constellationTranslated, style = 'GaNParagraph')
                
                if website1 in paragraph.text:
                    newLink = paragraph.text.replace("2018",str(year))
                    paragraph.text = None
                    paragraph.add_run(newLink, style = 'GaNLinks')
                
                elif website2 in paragraph.text:
                    newLink = paragraph.text.replace("2019",str(year))
                    paragraph.text = None
                    paragraph.add_run(newLink, style = 'GaNLinks')

    #Save a copy with a new name, date and language.
    dirPath = dirPath.rsplit('_', 2)[0]
    newWordPath = os.path.join(dirPath + "\\GaN_{year}_ActivityGuide_{cons}_lat_".format(year = year, cons = constName) + str(latitude) + "_" + str(languageBase) + ".docx")
    workingDoc.save(newWordPath)

    #Print information about the working file on
    print("The " + languageBase + " activity guide for the constellation {cons}".format(cons = constName) + " in the latitude {lat}".format(lat = latitude) +" south has been completed \n___________________________________________________________________________________________________________\n")

    # return the new doc path to make a list with it.
    return newWordPath



