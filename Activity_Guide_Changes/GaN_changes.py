# -*- coding:utf-8 -*-
# Install mtranslate, googletrans for translations
# Install python-docx for managing the Word Files.
# Install Pandas to manage the Excel file and bring the information
# Import Shutil to remove the directory

import os, time, sys 
from deep_translator import GoogleTranslator  
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.shared import RGBColor
import pandas as pd
from shutil import rmtree
import multiprocessing


def importNorthData():
    # Define the path for the excel file
    excelPath = os.path.join(sys.path[0],"GaN_cons_and_dates.xlsx")

    # Get Data from the Excel File using Pandas
    # Capitalize  constellations names for a later comparison
    dfNorth = pd.read_excel (excelPath,sheet_name='North')
    dfNorth['Constellations'] = dfNorth['Constellations'].str.capitalize()

    #store constellation and date information in respective variables
    northCons = dfNorth['Constellations']
    northDates = dfNorth['Dates']

    #updates the constellation: creates a variable to hold the new constellation, uses the old constellation
    #information to find the new constellation
    newConstellationNorth = {}
    
    # Ading key and values to the new dictionary
    for i in range(len(northCons)):
        newConstellationNorth[northCons[i]] = northDates[i] 

    #return the dictionary with the North Data
    return newConstellationNorth

def createDir(year, constellations):
    cons = constellations
    year = year
    savePath = os.getcwd() 
    savePath = os.path.join(savePath + "\GaN\docs_changed")
    rmtree(savePath)
    os.mkdir(savePath)
    paths = []
    for con in cons:
        savePath = os.getcwd() 
        savePath = os.path.join(savePath + "\GaN\docs_changed\GaN_{year}_ActivityGuide_{con}".format(year = year, con = con))        
        os.mkdir(savePath)
        paths.append(savePath)
    
    return paths

def createPaths(directories, languages):
    direcs = directories
    langs = languages

    dirPaths = []
    for lang in langs:
        for direc in direcs:
            dirPaths.append(direc + "_" + lang)
    return dirPaths





#opens document that will be edited
def openWordDoc(filename):
    document = Document(filename)
    return document


#updating Northern hemisphere information (constellation, date, text displayed to the user)
#######################################################################################
########################All the information that needs to change#######################
#######################################################################################

#######################################################################################
########################All the information that needs to change#######################
#######################################################################################

def northTranslation(dirPaths):
    #updating northern hemisphere information (constellation, date, text displayed to user)
    northConstellationReplacement = {
            
            "Catalan" : "Perseus",
            "Chinese" : "英仙座",
            "Czech" : "Persea",
            "English" : "Perseus",
            "Finnish" : "Perseus" ,
            "French" : "Persée" ,    
            "Galician" : "Perseo",
            "German" : "Perseus",
            "Greek" : "Περσεύς",
            "Indonesian" : "Perseus",
            "Japanese" : "ペルセウス",
            "Polish" : "Perseusz",
            "Portuguese" : "Perseu",
            "Romanian" : "Perseu",
            "Serbian" : "Персеус",
            "Slovak" : "Perseus",
            "Slovenian" : "Perseus",
            "Spanish" : "Perseo",
            "Swedish" : "Perseus",
            "Thai" : "เซอุส"
                
            }
        
    northDateReplacement = {
            
            "Catalan" : "30 d'octubre al novembre 8 i 29 de novembre de desembre 8",
            "Chinese" : "10月30日至11月 8月和11月29日至12月8",
            "Czech" : "30. října - 8. listopadu a 29. listopadu - 8. prosince",
            "English" : "Oct. 30-Nov. 8 and Nov. 29-Dec. 8",
            "Finnish" : "30 lokakuu- 8 marraskuu Ja 29 marraskuu-8 joulukuu" ,
            "French" : "Du 30 octobre au 8 novembre et du 29 novembre au 8 décembre" ,    
            "Galician" : "30 de outubro-8 de novembro e 29 de novembro-8 de decembro",
            "German" : "30. Oktober - 8. November und 29. November - 8. Dezember",
            "Greek" : "30 Οκτωβρίου-8 Νοεμβρίου και 29 Νοεμβρίου-8 Δεκεμβρίου",
            "Indonesian" : "30 Oktober-8 November dan 29 November-8 Desember",
            "Japanese" : "10月30日〜11月8日、11月29日〜12月8日",
            "Polish" : "30 października - 8 listopada i 29 listopada - 8 grudnia",
            "Portuguese" : "30 de outubro a 8 de novembro e 29 de novembro a 8 de dezembro",
            "Romanian" : "30 octombrie-8 noiembrie și 29 noiembrie-8 decembrie",
            "Serbian" : "30. октобра - 8. новембра и 29. новембра - 8. децембра",
            "Slovak" : "30. októbra - 8. novembra a 29. novembra - 8. decembra",
            "Slovenian" : "30. oktobra - 8. novembra in 29. novembra - 8. decembra",
            "Spanish" : "Del 30 de octubre al 8 de noviembre y del 29 de noviembre al 8 de diciembre",
            "Swedish" : "30. októbra - 8. novembra a 29. novembra - 8. decembra",
            "Thai" : "30 ตุลาคม - 8 พฤศจิกายนและ 29 พฤศจิกายน - 8 ธันวาคม"
                
            }

    northHeadingFirst = {
        "Catalan" : "Dates de la campanya ",
        "Chinese" : ""  ,
        "Czech" : "Informace v této příručce jsou určeny pro pozorovací kampaň probíhající od ",
        "English" : " ",
        "Finnish" : "",
        "French" : "Dates à utiliser pour la Campagne ",
        "Galician" : "Datas da campaña de ",
        "German" : "Kampagnendaten ",
        "Greek" : "",
        "Indonesian" : "Waktu Kampanye ",  
        "Japanese" : " ：",
        "Polish" : "",
        "Portuguese" : "Datas das campanhas de ",
        "Romanian" : "Perioadele campaniei din ",
        "Serbian" : "Сазвежђе ",
        "Slovak" : "V roku ",
        "Slovenian" :  "",
        "Spanish" :  "",
        "Swedish" : "Kampanjdatum för ",
        "Thai" : "กำหนดการในปีพ. ศ. "
        }

    northHeadingMiddle = {
        "Catalan" : " en què usem la  ",
        "Chinese" : "： "  ,
        "Czech" : ". Při pozorování použijte hvězdy oblohy, které zobrazují",
        "English" : " Campaign Dates that use ",
        "Finnish" : " havainnointijaksot vuonna ",
        "French" : " ",
        "Galician" : " que usan ",
        "German" : " für das ",
        "Greek" : " Ημερομηνίες παρατήρησης για τον  ",
        "Indonesian" : " untuk ",  
        "Japanese" : "年キャンペーン期間 (対象：",
        "Polish" : ": Daty kampanii używające ",
        "Portuguese" : " que usam ",
        "Romanian" : " pentru ",
        "Serbian" : " током ",
        "Slovak" : " môžete pozorovať ",
        "Slovenian" :  ": Datumi kampanje za opazovanje ",
        "Spanish" :  " Fechas de la campaña para ",
        "Swedish" : " ",
        "Thai" : " เซอุส"
        }

    northHeadingLast = {
        "Catalan" : " ",
        "Chinese" : "年"  ,
        "Czech" : ".",
        "English" : ": ",
        "Finnish" : ": ",
        "French" : ": ",
        "Galician" : ": ",
        "German" : ": ",
        "Greek" : ": ",
        "Indonesian" : ": ",  
        "Japanese" : ")：、",
        "Polish" : ": ",
        "Portuguese" : ": ",
        "Romanian" : ": ",
        "Serbian" : ". године посматрамо ",
        "Slovak" : ": ",
        "Slovenian" :  ": ",
        "Spanish" :  ": ",
        "Swedish" : ": ",
        "Thai" : "ดำเนินโครงการให้เสร็จสมบูรณ์: "
        }

    firstParagraphfirst = {
        "Chilean_Spanish" : "Usted está participando en una campaña mundial para observar y registrar las estrellas visibles más débiles como un medio para medir la contaminación lumínica en un lugar determinado. Localizando y observando la ",
        "Catalan" : "Esteu participant en una campanya mundial per observar i anotar la brillantor de les estrelles més febles que es poden veure, com a mitjà per mesurar la contaminació lumínica en un lloc determinat. Localitzant i observant la  ",
        "Chinese" : "你现在参加的是全球公益科普活动 Globe at Night （全球观星活动），这是一个以观察和记录夜空的可见恒星数来测量你所在地光污染情况的活动。通过定位和观测夜空中的",
        "Czech" : "Účastníte se celosvětové kampaně zaměřené na pozorování a záznam nejslabších viditelných hvězd jako prostředku měření světelného znečištění v daném místě. Lokalizací a pozorováním", ####Figure out what to do with the Czech one
        "English" : "You are participating in a global campaign to observe and record the faintest stars visible as a means of measuring light pollution in a given location. By locating and observing the constellation ",
        "Finnish" : "Osallistut maailmanlaajuiseen kampanjaan tarkkaillaksesi ja tallentaaksesi himmeimpiä näkyvissä olevia tähtiä keinona mitata valonsaastetta tietyssä paikassa. Paikallistamalla ja tarkkailemalla ",
        "French" :   "Vous allez participer à une campagne mondiale d’observation pour détecter les plus faibles étoiles visibles afin de mesurer la pollution lumineuse sur un site donné. Partout dans le monde, en localisant et en observant la ",
        "Galician" : "Grazas por participar nesta campaña global de medida da contaminación lumínica mediante a observación das estrelas máis febles que podes albiscar. Localizando e observando a  ",
        "German" : "Mach mit an einer weltweiten Kampagne, die schwächsten sichtbaren Sterne zu beobachten und aufzuzeichnen, um die Lichtverschmutzung an einem Ort zu messen. Durch das Auffinden und Beobachten des ",
        "Greek" : "Συμμετέχετε σε μία παγκόσμια καμπάνια για να παρατηρήσετε και να καταγράψετε τη φωτεινότητα των πιο αμυδρά ορατών άστρων σαν μέσο για την μέτρηση της Φωτορρύπανσης σε μία δεδομένη περιοχή. Με τον εντοπισμό και την παρατήρηση του  ",
        "Indonesian" : "Anda sedang berpartisipasi dalam kampanye global pengamatan dan pencatatan penampakan bintang paling redup untuk pengukuran tingkat polusi cahaya di suatu lokasi. Melalui pengamatan dan identifikasi  ",
        "Japanese" : '街には人工光があふれ、夜空が照らされ、星が見えにくくなってきています。また、無駄・過剰な人工光は、莫大なエネルギーの浪費、生態系への悪影響、人間生活・人体への悪影響をも引き起こしています。この光害（ひかりがい）の問題を啓発する活動に、あなたも参加してみませんか。Globe at Night（グローブ・アット・ナイト）は市民参加型の、夜空の明るさ世界同時観察キャンペーンです。どなたでも簡単に参加できます。決められた日時に屋外に出て夜空を眺め、星の見え方をインターネットで報告するだけ。ぜひあなたも参加して、光害の問題を考えてみませんか。そして、世界中の人と、美しい星空・地球環境への思いを共有しましょう。',
        "Polish" : "Uczestniczysz w ogólnoświatowym przedsięwzięciu, którego celem jest obserwacja i odnotowanie najsłabszych widocznych gwiazd w celu zmierzenia zanieczyszczenia światłem w danym miejscu. Poprzez zlokalizowanie i obserwację  ",
        "Portuguese" : "Está a participar numa campanha global para observar e registar as estrelas mais fracas visíveis como forma de medir a poluição luminosa num determinado local. Localizando e observando a  ",
        "Romanian" : "Prin această activitate participați în cadrul unei campanii globale de observare și consemnare a celor mai slabe stele vizibile ca metodă de măsurare a poluării luminoase dintr-un anumit loc. Localizând și observând  ",
        "Serbian" : "Ви сте учесници глобалног посматрачког пројекта, који има за циљ да одреди колико је светлосно загађене у средини у којој живите. Посматрајући звезде унутар  ",
        "Slovak" : "Stávate sa súčasťou celosvetovej kampane Globe at Night, ktorej cieľom je meranie svetelného znečistenia. Pozorovaním  ",
        "Slovenian" : "Sodelujete v svetovni aktivnosti opazovanja in beleženja najšibkejših, s prostim očesom  še vidnih zvezd, kot metode za merjenje svetlobnega onesnaževanja na določenem mestu. Z opazovanjem izbranega ",
        "Spanish" : "Usted está participando en una campaña mundial para observar y registrar las estrellas visibles más débiles como un medio para medir la contaminación lumínica en un lugar determinado. Localizando y observando la  ",
        "Swedish" : "Du deltar i en världsomspännande kampanj för att observera och rapportera de svagaste synliga stjärnorna, som ett mått på ljusföroreningarna på orten. Genom att hitta och observera ",
        "Thai" : "คุณกำลังร่วมนโครงการระดับโลกที่จะสังเกตและบันทึกผลดาวฤกษ์ที่จางที่สุดที่มองเห็นได้ ซึ่งก็คือการวัดมลพิษทางแสงในสถานที่นั้นๆ  โดยการมองหาและสังเกต "

    }

    firstParagraphLast = {
        "Chilean_Spanish" : " en el cielo nocturno y comparándolo con las cartas estelares, la gente de todo el mundo aprenderá cómo las luces de su comunidad contribuyen a la contaminación lumínica. Sus contribuciones a la base de datos en línea documentarán el cielo nocturno visible.",
        "Catalan" : " a la nit i comparant la brillantor de les estrelles del cel amb la brillantor que indiquen els mapes, gent de tot el món aprendran com els llums de la seva zona contribueixen a augmentar la contaminació lumínica. Les vostres aportacions a la base de dades activa faran palesa la visibilitat del cel nocturn.",
        "Chinese" : "，并将所肉眼观察到的星等情况与所给出的星等图表作对比，我们可以知道自己社区中的人造光是如何导致光污染的。你所提供数据将和来自全世界的数据一起帮助建立一张全球光污染地图。",
        "Czech" : "na noční obloze a jejím porovnáním s hvězdnými mapami se lidé z celého světa dozvědí, jak světla v jejich komunitě přispívají ke světelnému znečištění. Vaše příspěvky do online databáze budou dokumentovat viditelnou noční oblohu.", ###Figure something out with the Czech
        "English" : " in the night sky and comparing it to stellar charts, people from around the world will learn how the lights in their community contribute to light pollution. Your contributions to the online database will document the visible nighttime sky.",
        "Finnish" : " miten valosaaste syntyy kunkin taajaman tai muun ihmisen toiminnan valoista. Antamasi tiedot päivittyvät heti verkossa olevaan tietokantaan, ja näin saadaan käsitys siitä minkä verran taivaan tähdistä on missäkin nähtävissä.",
        "French" :   " dans le ciel nocturne et en la comparant aux cartes stellaires, les participants, apprendront comment l’éclairage, dans leur environnement local, influence la pollution lumineuse. Vos contributions à la base de données en ligne permettront de mesurer la qualité du ciel nocturne.",
        "Galician" : " e comparándoa co que aparece nos mapas estelares recollidos neste documento podes saber canto contribúen á contaminación lumínica os sistemas de iluminación que hai no teu barrio ou vila. As túas achegas á base de datos en liña de GLOBE at Night (O MUNDO á Noite) servirán para documentar a calidade do ceo nocturno.",
        "German" : " am Nachthimmel und den Vergleich mit den Helligkeitskarten, lernen Menschen auf der ganzen Erde, wie die Lichter in ihrer Gemeinde zur Lichtverschmutzung beitragen. Dein Beitrag zur Online-Datenbank beschreibt den sichtbaren Nachthimmel.",
        "Greek" : " στον νυχτερινό ουρανό καθώς και με την σύγκριση των ανωτέρω με τα διαγράμματα για τα μεγέθη των άστρων,  άνθρωποι από όλον τον κόσμο θα μάθουν πώς τα φώτα στην κοινότητά τους συμβάλλουν στην Φωτορρύπανση. Με την κατάθεση των πορισμάτων τους στην ιστοσελίδα θα δημιουργηθεί ένα αρχείο σχετικά με το τι μπορεί να δει κανείς στον νυχτερινό ουρανό.",
        "Indonesian" : " di langit malam dan membandingkannya dengan peta bintang, masyarakat di seluruh dunia dapat mengetahui dan mempelajari seberapa besar kontribusi cahaya di lingkungannya terhadap polusi cahaya. Kontribusi data anda pada basis data online akan membantu mendokumentasikan langit malam yang tampak di berbagai lokasi.",
        "Japanese" : "",
        "Polish" : " na nocnym niebie oraz porównanie go do map nieba ludzie z całego świata będą mogli dowiedzieć się jaki wkład światło emitowane przez ich społeczność wnosi do  zanieczyszczenia światłem. To co dodasz do internetowej bazy danych pomoże udokumentować widoczne nocne niebo.",
        "Portuguese" : " no céu noturno e,  comparando-a com cartas estelares, pessoas de todo o mundo aprenderão  como as luzes da sua comunidade contribuem para a poluição luminosa. As suas contribuições para a base de dados on-line irão documentar a visibilidade do céu noturno em todo o mundo.",
        "Romanian" : " pe cerul nopții și comparând-o cu diagramele stelare, oamenii din întreaga lume vor putea afla în ce măsură iluminatul nocturn din comunitatea lor contribuie la poluarea luminoasă. Contribuțiile dumneavoastră la baza de date online vor facilita o documentare globală privind cerul nocturn observabil.",
        "Serbian" : " и упоређујући их са приложеним звезданим картама, посматрачи широм света могу на практичном примеру да увиде колико је светлосно загађење у њиховој средини. Кроз учешће у овом пројекту, допринећете целовитијем сагледавању глобалног проблема.",
        "Slovak" : " na nočnej oblohe a porovnávaním skutočnej situácie s našimi mapkami sa nielenže dozviete, ako osvetlenie vo Vašom okolí prispieva k svetelnému znečisteniu, ale budete môcť porovnať úroveň svetelného znečistenia aj s inými lokalitami z celého sveta. Vaše pozorovanie tiež rozšíri online databázu dokumentujúcu viditeľnosť nočnej oblohy na našej planéte",
        "Slovenian" : " na nočnem nebu in s primerjavo videnega z zvezdnimi kartami, se lahko ljudje širom sveta podučijo o tem, kako svetila v njihovem kraju prispevajo k svetlobnemu onesnaževanju.  Vaši prispevki v spletno bazo podatkov bodo pomagali dokumentirati nočno nebo, vidno s prostim očesom.",
        "Spanish" : " en el cielo nocturno y comparándolo con las cartas estelares, la gente de todo el mundo aprenderán cómo las luces de su comunidad contribuyen a la contaminación lumínica. Sus contribuciones a la base de datos en línea documentarán el cielo nocturno visible.",
        "Swedish" : " på natthimlen kan folk i hela världen lära sig hur belysningen i våra samhällen och omgivningar bidrar till ljusföroreningar. Era bidrag till online-databasen hjälper till att dokumentera den synliga natthimlens över hela världen.",
        "Thai" : "ในท้องฟ้ายามค่ำคืนและเปรียบเทียบสิ่งที่เห็นกับแผนภาพที่เราให้า คนจากทั่วทุกมุมโลกจะได้เรียนรู้ว่าแสงไฟในชุมชนของพวกเขาสร้างมลพิษทางแสงอย่างไร ผลงานของคุณจะอยู่ในถูกเก็บในฐานข้อมูลออนไลน์ ซึ่งจะเป็นเอกสารเกี่ยวกับท้องฟ้ายามค่ำคืนที่เรามองเห็น",

    }

    ##################################################################################################
    ##################################################################################################
            ###	End of the changes section defining things that need to be changed			###
    ##################################################################################################
    ##################################################################################################

    # Get data from the Excel file and bring the created Paths
    dirPath = dirPaths
    northData = importNorthData()


    #1date2con3
    CountryList1 = ("Czech")
    #1con2year3date
    CountryList2 = ("chinese (traditional)", "Finnish", "Serbian", "Swedish")
    #1year2Con3date
    CountryList3 = ("Chilean_Spanish", "Catalan", "English", "French", "Galician", "German", "Greek", "Indonesian", "Japanese", "Polish", "Portuguese", "Romanian", "Slovak", "Slovenian", "Spanish", "Thai")  #1year2Con3date

    # Getting data from the Paths
    languageBase = dirPath.split('_')[-1]
    constName = dirPath.split('_')[-2]
    year = dirPath.split('_')[-4]
    thaiYear = int(year)+ 543

    #Be sure to change the websites into the word files
    website1 = "astro/maps/GaNight/2018/"
    website2 = "astro/maps/GaNight/2019/"

    
    # Define the Word file path as the original file
    wordPath = os.path.abspath("..\Gan\GaN\docs_to_change\GaN2018_ActivityGuide_Perseus_N_")
    workingDoc = openWordDoc(wordPath + str(languageBase) + ".docx")

    # Chinese in not in the dictionary of deep_translator, is better "chinese (traditional)"
    if languageBase != "Chinese":
        languageBase = languageBase
    else:
        languageBase = "chinese (traditional)"

    # styles of each paragraph to kkep the original word styles
    objStyles = workingDoc.styles
    objCharstyle = objStyles.add_style('GaNStyle', WD_STYLE_TYPE.CHARACTER)
    objFont = objCharstyle.font
    objFont.name = 'Calibri'
    objFont.size = Pt(14)
    
    objStyles2 = workingDoc.styles
    objCharstyle2 = objStyles2.add_style('GaNParagraph', WD_STYLE_TYPE.CHARACTER)
    objFont2 = objCharstyle2.font
    objFont2.name = 'Calibri'
    objFont2.size = Pt(10)
    
    objStyles3 = workingDoc.styles
    objCharstyle3 = objStyles3.add_style('GaNLinks', WD_STYLE_TYPE.CHARACTER)
    objFont3 = objCharstyle3.font
    objFont3.name = 'Calibri'
    objFont3.size = Pt(9.5)
    objFont3.bold = True
    objFont3.underline = True
    objFont3.color.rgb = RGBColor(0,0,128)


    #Define the base language in deep_translator and translate it into de destiny language
    constellationTranslated =GoogleTranslator(source ='english', target = languageBase.lower()).translate(constName +" constellation")
    dateTranslated = GoogleTranslator(source ='english', target = languageBase.lower()).translate(northData.get(constName))
    print(languageBase.lower())

    # Replace the translations in the proper places
    for languageSelected, date in northDateReplacement.items():
        if languageSelected == languageBase:
            for paragraph in workingDoc.paragraphs:
                #If the contellation's name is in the paragraph, delete the paragraph and add a new one with the translations
                if northConstellationReplacement[languageBase] in paragraph.text:
                    # Replace only if the name and the date is on the paragraph, organizng with the grammar of each language
                    if date in paragraph.text:
                        paragraph.clear()
                        if languageBase in CountryList1:
                            paragraph.add_run(northHeadingFirst[languageBase]+ dateTranslated +northHeadingMiddle[languageBase]+ constellationTranslated + northHeadingLast[languageBase], style = 'GaNStyle')
                        elif languageBase in CountryList2:
                            paragraph.add_run(northHeadingFirst[languageBase]+ constellationTranslated + northHeadingMiddle[languageBase]+ str(year) +northHeadingLast[languageBase]+ dateTranslated + ".", style = 'GaNStyle' )
                        elif languageBase in CountryList3:
                            if languageBase != "Thai":      
                                paragraph.add_run(northHeadingFirst[languageBase]+ str(year) +northHeadingMiddle[languageBase] + constellationTranslated +northHeadingLast[languageBase] + dateTranslated, style = 'GaNStyle')
                            else:
                                paragraph.add_run(northHeadingFirst[languageBase]+ str(thaiYear) +northHeadingMiddle[languageBase] + constellationTranslated +northHeadingLast[languageBase] + dateTranslated, style = 'GaNStyle')
                    # Replace only if the constellation's name is in the paragraph
                    else:
                        paragraph.clear()
                        if(languageBase!= 'Japanese'):
                            paragraph.add_run(firstParagraphfirst[languageBase] + constellationTranslated + firstParagraphLast[languageBase], style = 'GaNParagraph')
                        else:
                            paragraph.add_run(firstParagraphfirst[languageBase] + firstParagraphLast[languageBase] + constellationTranslated, style = 'GaNParagraph')
                
                if website1 in paragraph.text:
                    newLink = paragraph.text.replace("2018",str(year))
                    paragraph.text = None
                    paragraph.add_run(newLink, style = 'GaNLinks')
                
                elif website2 in paragraph.text:
                    newLink = paragraph.text.replace("2019",str(year))
                    paragraph.text = None
                    paragraph.add_run(newLink, style = 'GaNLinks')

    #Save a copy with a new name, date and language.
    dirPath = dirPath.rsplit('_', 1)[0]
    newWordPath = os.path.join(dirPath + "\GaN_{year}_ActivityGuide_{cons}_".format(year = year, cons = constName) + str(languageBase) + ".docx")
    workingDoc.save(newWordPath)

    #Print information about the working file on
    return print("The " + languageBase + " activity guide for the constellation {cons}".format(cons = constName) + " has been completed \n____________________________________________________________________________________________\n")



if __name__ =='__main__':

    # Start time counter
    start = time.time()

    # Get the data from the User
    year = 2022
    constellations = ["Perseus", "Taurus", "Gemini", "Leo", "Bootes", "Cygnus", "Pegasus", "Orion", "Hercules"]
    languages = ["Catalan", "Chinese", "Czech", "English", "Finnish", "French", "Galician", "German", "Greek", "Indonesian", "Japanese", "Polish", "Portuguese", "Romanian", "Serbian", "Slovak", "Slovenian", "Spanish", "Swedish", "Thai"]
    
    # Creating the directories and the Paths
    directories= createDir(year, constellations)
    paths = createPaths(directories, languages)
    
    #Calll de translation function, requiring multiprocessing with Pool
    pool = multiprocessing.Pool(processes = 4)
    for path in paths:
        pool.apply_async(northTranslation, args = (path, ))
    pool.close()
    pool.join()


    # Finishing time counter and getting time of execution
    finish = time.time() - start
    print('Execution time: ', time.strftime("%H:%M:%S", time.gmtime(finish)))
