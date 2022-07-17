from docx import Document
from PIL import Image
from docx.shared import Pt, Inches
import os

# Open the MS Word files created in the translations
def openWordDoc(filename):
    document = Document(filename)
    return document

# Get the latitudes according to the images names in the GaN webpage (north ="10", south ="10s")
def transformLatitude(lat):
    if "N" in lat:
        lat = str(lat.rstrip(lat[-1]))
    else:
        lat = str(lat.lower())
    return lat

# get the links from the charts for each latitude
def openImage(fileName):
    constellation = fileName.split('_')[-4]
    latitude = fileName.split('_')[-2]

    path = os.getcwd() 
    path = os.path.join(path + "\GaN\images")

    lat = transformLatitude(latitude)

    magnitudes = ["05", "15", "25", "35", "45", "55", "65", "75"]

    pathsList = []
    for mag in magnitudes:
        pathsList.append(path + "\\" + constellation + "-" + lat + "_" + mag + ".png")
    
    return pathsList




#crops an image based on the magnitude, saves image as png
def cutStarChart(filename):

    workingDoc = openWordDoc(filename)
    dirCharts = openImage(filename)

    for dc in dirCharts:
        workingDoc.add_picture(dc)
        print(dc)

    
    
    '''
    cutPoints = {
            "Mag0" : (80,80,826,596),
            "Mag1" : (834,80,1577,596),
            "Mag2" : (80,604,826,1119),
            "Mag3" : (834,604,1577,1119),
            "Mag4" : (80,1127,826,1642),
            "Mag5" : (834,1127,1577,1642),
            "Mag6" : (80,1651,826,2166),
            "Mag7" : (834,1651,1577,2166)
            }

    for key, value in cutPoints.items():
        im = Image.open(filename)
        im.crop(value).save(key + ".png")

        i = 0
        k = 0
        l = 0              

        for table in workingDoc.tables:
            
            print(table.cell)
            Big_picture_cells = (table.cell(1,0), table.cell(1,2), table.cell(4,0), table.cell(4,2))
            small_picture_cells = (table.cell(1,0), table.cell(1,1), table.cell(1,2), table.cell(1,3), table.cell(3,0), table.cell(3,1), table.cell(3,2), table.cell(3,3))
            
            if i < 2:
                for tableCell in Big_picture_cells:
                    tableCell.paragraphs[1].clear()
                                        
                    tableCell.paragraphs[1].add_run().add_picture("Mag" + str(k) + ".png", width = Inches(3.39), height = Inches(2.35))
                    workingDoc.save(filename)
                    k = k + 1
            
            else :
                for tableCell in small_picture_cells:
                    tableCell.paragraphs[0].clear()
                    tableCell.paragraphs[0].add_run().add_picture("Mag" + str(l) + ".png", width = Inches(1.44), height = Inches(1.01))
                    workingDoc.save(filename)
                    l = l + 1
                
            i = i + 1  
        '''

    workingDoc.save(filename)   

def bigTable():
    pass

def littleTable():
    pass

northListPaths = ['C:\\Users\\Marco Moreno\\OneDrive\\Documentos\\Enciso Systems\\GaN\\GaN\\docs_changed\\GaN_North_2022_ActivityGuide_Leo\\GaN_2022_ActivityGuide_Leo_lat_0_Spanish.docx', 'C:\\Users\\Marco Moreno\\OneDrive\\Documentos\\Enciso Systems\\GaN\\GaN\\docs_changed\\GaN_North_2022_ActivityGuide_Leo\\GaN_2022_ActivityGuide_Leo_lat_10N_Spanish.docx']

southListPaths = ['C:\\Users\\Marco Moreno\\OneDrive\\Documentos\\Enciso Systems\\GaN\\GaN\\docs_changed\\GaN_South_2022_ActivityGuide_Orion\\GaN_2022_ActivityGuide_Orion_lat_10S_French.docx', 'C:\\Users\\Marco Moreno\\OneDrive\\Documentos\\Enciso Systems\\GaN\\GaN\\docs_changed\\GaN_South_2022_ActivityGuide_Orion\\GaN_2022_ActivityGuide_Orion_lat_20S_French.docx']



path1 = northListPaths[0]
path2 = southListPaths[0]
print(path1)
cutStarChart(path1)


