from docx import Document
import os

def openWordDoc(filename):
    document = Document(filename)
    return document

workingDoc = openWordDoc(r'C:\Users\Marco Moreno\OneDrive\Documentos\Enciso Systems\GaN\GaN\docs_to_change\GaN2018_ActivityGuide_Perseus_N_Catalan.docx')

lineSearched = "30 d'octubre al novembre 8 i 29 de novembre de desembre 8"

for paragraph in workingDoc.paragraphs:
    if lineSearched in paragraph.text:
        paragraph.clear()
        paragraph.add_run('esto es lo que busco')

newWordPath = r"C:\Users\Marco Moreno\OneDrive\Documentos\Enciso Systems\GaN\GaN\docs_changed\este.docx"
workingDoc.save(newWordPath)
